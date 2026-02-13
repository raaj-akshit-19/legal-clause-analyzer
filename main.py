import io
import json
import tempfile
from fastapi import FastAPI, Form, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from analyzer import analyze_clauses
from fpdf import FPDF
from docx import Document
import os

app = FastAPI()

app.mount("/static", StaticFiles(directory="static"), name="static")

# Allow CORS for development convenience
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def read_index():
    return FileResponse('static/index.html')

@app.post("/analyze")
async def analyze(text: str = Form(...)):
    if not text:
        raise HTTPException(status_code=400, detail="Text is required")
    results = analyze_clauses(text)
    return JSONResponse(content={"results": results})

@app.post("/download")
async def download_file(format: str = Form(...), data: str = Form(...)):
    try:
        clauses = json.loads(data)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON data")

    if format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Legal Clause Analysis", ln=True, align='C')
        pdf.ln(10)
        
        for clause in clauses:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, txt=f"Type: {clause['type']}", ln=True)
            
            pdf.set_font("Arial", size=12)
            # Multi_cell covers text wrapping
            pdf.multi_cell(0, 10, txt=f"Condition: {clause['condition']}")
            pdf.multi_cell(0, 10, txt=f"Consequence: {clause['consequence']}")
            pdf.ln(5)
            
        # Output to a temporary file
        filename = "analysis_report.pdf"
        output_path = os.path.join(tempfile.gettempdir(), filename)
        pdf.output(output_path)
        
        return FileResponse(output_path, media_type='application/pdf', filename=filename)

    elif format == "docx":
        doc = Document()
        doc.add_heading('Legal Clause Analysis', 0)

        for clause in clauses:
            doc.add_heading(clause['type'], level=1)
            doc.add_paragraph(f"Condition: {clause['condition']}")
            doc.add_paragraph(f"Consequence: {clause['consequence']}")
            doc.add_paragraph("") # Space

        filename = "analysis_report.docx"
        output_path = os.path.join(tempfile.gettempdir(), filename)
        doc.save(output_path)
        
        return FileResponse(output_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=filename)

    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use 'pdf' or 'docx'.")
